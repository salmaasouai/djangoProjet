from django import forms  # Add this import
from django.views import generic
from django.urls import reverse_lazy
from django.shortcuts import render, get_object_or_404, redirect
from .models import Event, Participant
from .forms import ParticipantForm  # Ensure this form is defined
from django.contrib import messages
import csv
from django.http import HttpResponse
from django.template.loader import render_to_string
import pandas as pd
import io
import xhtml2pdf.pisa as pisa
from docx import Document
from django.template.loader import get_template  # Importation de get_template
from django.db.models import Avg
from users.models import Client  # Import your Client model

import matplotlib.pyplot as plt
import io
import base64
from django.db.models import Count


from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.contrib import messages
import http.client
import json
from .forms import EmailForm

# views.py

from django.db.models import Q
from .models import Event


from django.views import generic
from .models import Event



from django.db.models import Q

class EventListView(generic.ListView):
    model = Event
    template_name = 'events/event_list.html'
    context_object_name = 'events'

    def get_queryset(self):
        queryset = Event.objects.all()  # Default queryset

        # Access the request object via self.request
        search_query = self.request.GET.get('q', '')  # Récupère la valeur de la recherche
        if search_query:
            # Utilisation de Q pour combiner les conditions de recherche avec un OR
            queryset = Event.objects.filter(
                Q(title__icontains=search_query) |
                Q(description__icontains=search_query) |
                Q(date__icontains=search_query) |
                Q(location__icontains=search_query) |
                Q(organizer__icontains=search_query) |
                Q(price__icontains=search_query) |
                Q(created_at__icontains=search_query) |
                Q(updated_at__icontains=search_query)
            )

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        

        # Calculate the remaining capacity for each event
        for event in context['events']:
            event.capacity_remaining = event.max_participants - event.participants.count()

        return context




class EventCreateForm(forms.ModelForm):
    class Meta:
        model = Event
        fields = [
            'title', 
            'description', 
            'date', 
            'location', 
            'organizer',
            'max_participants', 
            'price', 
            'image'  
        ]

class EventCreateView(generic.CreateView):
    model = Event
    template_name = 'events/event_form.html'
    form_class = EventCreateForm
    success_url = reverse_lazy('event_list')

    def form_valid(self, form):
        return super().form_valid(form)

class EventDetailView(generic.DetailView):
    model = Event
    template_name = 'events/event_detail.html'

class EventUpdateView(generic.UpdateView):
    model = Event
    template_name = 'events/event_form.html'
    fields = [
        'title', 
        'description', 
        'date', 
        'location', 
        'organizer', 
        'max_participants', 
        'price', 
        'image', 
        'organizer'
        ]  
    success_url = reverse_lazy('event_list')

    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        
        # Pre-fill the date with the current value and make it read-only
        form.fields['date'].initial = self.object.date
        form.fields['date'].widget.attrs['readonly'] = True
        return form

    def form_valid(self, form):
        # Always keep the original date
        form.instance.date = self.object.date

        # Keep the original image if no new image is provided
        if not form.cleaned_data.get('image'):
            form.instance.image = self.object.image
        
        return super().form_valid(form)



class EventDeleteView(generic.DeleteView):
    model = Event
    template_name = 'events/event_confirm_delete.html'
    success_url = reverse_lazy('event_list')


 
def participant_list(request, event_id):
    event = get_object_or_404(Event, id=event_id)
    participants = Participant.objects.filter(event_id=event_id)
    
    # Filtrage par niveau de fitness
    fitness_level = request.GET.get('fitness_level')
    if fitness_level:
        participants = participants.filter(fitness_level=fitness_level)

    # Calcul de la capacité restante
    capacity_remaining = event.max_participants - participants.count()

    # Passer les participants triés et filtrés à la vue
    return render(request, 'participant/participant_list.html', {
        'event': event,
        'participants': participants,
        'capacity_remaining': capacity_remaining,
    })



def participant_create(request, event_id):
    # Récupérer tous les clients
    clients = Client.objects.all()

    # Récupération de l'événement
    event = get_object_or_404(Event, id=event_id)

    # Récupérer les participants existants à cet événement
    participants = Participant.objects.filter(event=event)

    # Calcul de la capacité restante
    capacity_remaining = event.max_participants - participants.count()

    # Vérification si la capacité est atteinte
    if capacity_remaining <= 0:
        messages.error(request, "La capacité maximale a été atteinte.")
        return redirect('participant_list', event_id=event.id)

    # Traitement du formulaire en POST
    if request.method == 'POST':
        form = ParticipantForm(request.POST)

        if form.is_valid():
            # Récupération du client à partir des données du formulaire
            client = form.cleaned_data['client']

            # Vérification si le client a déjà participé à cet événement
            if participants.filter(client=client).exists():
                messages.error(request, "Ce client a déjà participé à cet événement.")
                return redirect('participant_list', event_id=event.id)

            # Enregistrement du participant avec l'événement lié
            participant = form.save(commit=False)
            participant.event = event

            # Définir l'icône de statut de paiement
            if participant.payment_status:
                participant.payment_status_icon = 'tick-green'  # Assurez-vous d'utiliser cette icône dans votre template
            else:
                participant.payment_status_icon = 'tick-red'

            participant.save()

            messages.success(request, "Participant ajouté avec succès.")
            return redirect('participant_list', event_id=event.id)
        else:
            messages.error(request, "Veuillez corriger les erreurs dans le formulaire.")
    else:
        form = ParticipantForm()

    # Affichage de la page de création de participant
    return render(request, 'participant/participant_create.html', {
        'event': event,
        'form': form,
        'capacity_remaining': capacity_remaining,
        'clients': clients,
    })




def participant_update(request, event_id, participant_id):

    event = get_object_or_404(Event, id=event_id)  # Fetch the event based on the ID
    participant = get_object_or_404(Participant, id=participant_id)  # Fetch the participant based on the ID

    if request.method == 'POST':
        # Handle form submission and saving
        # For example, using a form to handle participant updates:
        form = ParticipantForm(request.POST, instance=participant)
        if form.is_valid():
            form.save()
            return redirect('participant_list', event_id=event.id)  # Redirect to the participant list for the event
    else:
        form = ParticipantForm(instance=participant)

    return render(request, 'participant/participant_update.html', {
        'form': form,
        'participant': participant,
        'event': event,  # Pass the event to the template context
    })




def participant_delete(request, event_id, participant_id):
    participant = get_object_or_404(Participant, pk=participant_id)
    if request.method == 'POST':
        participant.delete()
        return redirect('participant_list', event_id=event_id)
    return render(request, 'participant/participant_confirm_delete.html', {'participant': participant})



def event_export(request, format):
    events = Event.objects.all()

    # Créer un tableau pour stocker les événements et leurs participants restants
    for event in events:
        event.capacity_remaining = event.max_participants - event.participants.count()

    if format == 'pdf':
        # Générer un PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=events.pdf'

        # Code de génération de PDF
        template_path = 'events/event_pdf_template.html'  # Chemin vers votre template PDF
        template = get_template(template_path)
        context = {'events': events}
        html = template.render(context)

        pisa_status = pisa.CreatePDF(html, dest=response)

        if pisa_status.err:
            return HttpResponse('Une erreur est survenue lors de la génération du PDF.')

        return response

    elif format == 'docs':
        # Générer un fichier Docs
        response = HttpResponse(content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=events.doc'

        doc = Document()
        doc.add_heading('Liste des Événements', level=1)

        for event in events:
            doc.add_heading(event.title, level=2)
            doc.add_paragraph(f'Description: {event.description}')
            doc.add_paragraph(f'Date: {event.date.strftime("%d/%m/%Y %H:%M")}')
            doc.add_paragraph(f'Lieu: {event.location}')
            doc.add_paragraph(f'Organisateur: {event.organizer}')
            doc.add_paragraph(f'Participants Maximum: {event.max_participants}')
            doc.add_paragraph(f'Participants Restants: {event.capacity_remaining}')

            doc.add_paragraph(f'Prix: {event.price} €')
            doc.add_paragraph(f'Date de création: {event.created_at}')
            doc.add_paragraph(f'Date de Modification: {event.updated_at}')

            doc.add_paragraph('---')  # Ligne de séparation entre les événements

        doc.save(response)  # Enregistrer le document dans la réponse
        return response

    elif format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="events.csv"'

        writer = csv.writer(response)
        writer.writerow(['Titre', 'Description', 'Date', 'Lieu', 'Organisateur', 'Participants Maximum', 'Prix', 'Created_at', 'Updated_at'])  # En-tête

        for event in events:
            writer.writerow([
                event.title,
                event.description,
                event.date.strftime('%d/%m/%Y %H:%M'),  # Formatage de la date
                event.location,
                event.organizer,
                event.max_participants,
                event.price,
                event.created_at.strftime('%d/%m/%Y %H:%M'),  # Formatage de la date de création
                event.updated_at.strftime('%d/%m/%Y %H:%M')  # Formatage de la date de modification
            ])

        return response

    else:
        # Gérer les formats non pris en charge
        return HttpResponse("Format non pris en charge", status=400)





def send_email(request, event_id):
    event = get_object_or_404(Event, id=event_id)
    clients = Client.objects.all()  # Récupère tous les clients disponibles pour le formulaire

    if request.method == "POST":
        form = EmailForm(request.POST)
        if form.is_valid():
            # Récupération du client et du message à partir des données du formulaire
            client = form.cleaned_data['client']  # Le client est récupéré par email
            title = form.cleaned_data['title']  # Le message du formulaire
            message = form.cleaned_data['message']  # Le message du formulaire

            # Connexion à l'API pour envoyer l'email
            conn = http.client.HTTPSConnection("rapidmail.p.rapidapi.com")
            payload = json.dumps({
                "sendto": client.email,  # L'email du client depuis la base de données
                "name": f"{client.first_name} {client.last_name}",
                "replyTo": "raifguizani10@gmail.com",
                "ishtml": "false",
                "title": title,  # Titre du message, à personnaliser
                "body": message  # Le message que vous avez saisi dans le formulaire
            })

            headers = {
                'x-rapidapi-key': "f875700bacmshc39a326dde2ffc3p163701jsn487deac76591",
                'x-rapidapi-host': "rapidmail.p.rapidapi.com",
                'Content-Type': "application/json"
            }

            # Envoi de la requête
            conn.request("POST", "/", payload, headers)
            res = conn.getresponse()
            data = res.read()
            response_data = data.decode("utf-8")

            if "success" in response_data.lower():
                messages.success(request, "Email sent successfully!")  # Message de succès
                return redirect('participant_list', event_id=event_id)  # Redirection vers la page des participants après envoi
            else:
                messages.error(request, "Failed to send email.")  # Message d'erreur
                return redirect('participant_list', event_id=event_id)  # Redirection en cas d'erreur
    else:
        form = EmailForm()

    return render(request, 'participant/send-email.html', {'form': form, 'event': event, 'clients': clients})








def event_statistique(request, event_id):
    event = Event.objects.get(id=event_id)  # Get the event by its ID

    # Pie chart: Distribution of participants across events
    events = Event.objects.all()
    participants_count = [event.participants.count() for event in events]
    labels = [event.title for event in events]

    plt.figure(figsize=(6, 6))
    plt.pie(participants_count, labels=labels, autopct='%1.1f%%')
    plt.title("Participants Distribution by Event")
    pie_buffer = io.BytesIO()
    plt.savefig(pie_buffer, format='png')
    pie_buffer.seek(0)
    pie_base64 = base64.b64encode(pie_buffer.read()).decode()
    pie_buffer.close()

    # Line chart: Number of participants over time
    events = Event.objects.order_by('date')
    dates = [event.date.strftime("%Y-%m-%d") for event in events]
    participants_count = [event.participants.count() for event in events]

    plt.figure(figsize=(10, 5))
    plt.plot(dates, participants_count, marker='o')
    plt.xticks(rotation=45)
    plt.title("Number of Participants Over Time")
    plt.xlabel("Event Date")
    plt.ylabel("Number of Participants")
    line_buffer = io.BytesIO()
    plt.savefig(line_buffer, format='png')
    line_buffer.seek(0)
    line_base64 = base64.b64encode(line_buffer.read()).decode()
    line_buffer.close()

    return render(request, 'events/event_statistique.html', {
        'pie_chart': pie_base64,
        'line_chart': line_base64,
        'event': event,
    })






import qrcode
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from io import BytesIO
import base64
from django.templatetags.static import static
from .models import Event

def generate_qr_code(request, event_id):
    # Récupérer l'événement à partir de l'ID
    event = get_object_or_404(Event, pk=event_id)
    event.capacity_remaining = event.max_participants - event.participants.count()


 # Vérifier si l'événement a une image
    if event.image:
        # Lire l'image de l'événement
        image_file = event.image
        img_data = image_file.read()
    else:
        # Utiliser une image par défaut si l'événement n'a pas d'image
        default_image_path = static('images/default-event.jpg')
        with open(default_image_path, "rb") as f:
            img_data = f.read()
            
    # Détails de l'événement à inclure dans le QR code
    event_details = {
        "title": event.title,
        "image_file": image_file,
        "description": event.description,
        "date": event.date.strftime('%Y-%m-%d %H:%M'),
        "location": event.location,
        "price": f"{event.price} TND",
        "max_participants": event.max_participants,
        "remaining_participants": event.capacity_remaining,
        "created_at": event.created_at.strftime('%Y-%m-%d %H:%M'),
        "updated_at": event.updated_at.strftime('%Y-%m-%d %H:%M'),
    }

    # Convertir les détails de l'événement en une chaîne
    event_details_str = f"Title: {event_details['title']}\n" \
                        f"Image (Base64): {event_details['image_file']}\n" \
                        f"Description: {event_details['description']}\n" \
                        f"Date: {event_details['date']}\n" \
                        f"Location: {event_details['location']}\n" \
                        f"Price: {event_details['price']}\n" \
                        f"Max Participants: {event_details['max_participants']}\n" \
                        f"Remaining Participants: {event_details['remaining_participants']}\n" \
                        f"Created At: {event_details['created_at']}\n" \
                        f"Updated At: {event_details['updated_at']}"

    # Créer le QR code avec la version spécifiée
    qr = qrcode.QRCode(
        version=4,  # Spécifiez une version dans la plage de 1 à 40
        error_correction=qrcode.constants.ERROR_CORRECT_L,  # Niveau de correction d'erreur
        box_size=5,  # Taille de chaque box dans le QR code
        border=4,  # Bordure autour du QR code
    )
    qr.add_data(event_details_str)
    qr.make(fit=True)

    # Convertir le QR code en image binaire
    img_io = BytesIO()
    qr.make_image(fill='black', back_color='white').save(img_io, 'PNG')
    img_io.seek(0)

    # Retourner l'image QR code dans la réponse HTTP
    response = HttpResponse(img_io, content_type='image/png')
    return response
